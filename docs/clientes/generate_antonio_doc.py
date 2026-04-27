"""Genera el documento Word de respuesta para Antonio (J/29 1987)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


AZUL = RGBColor(0x0B, 0x3D, 0x91)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x00, 0x00, 0x00)


def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AZUL
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, size=11, color=NEGRO, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    return p


def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Encabezado
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("IMPORLAN")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = AZUL
    r.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Importación de Embarcaciones desde USA a Chile")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GRIS
    r.font.name = "Calibri"

    doc.add_paragraph()

    # Datos del documento
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = info.add_run("Fecha: 27 de abril de 2026\nPara: Antonio Pi\nAsunto: Cotización y proceso de importación – J/29 1987")
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS
    r.font.name = "Calibri"

    add_heading(doc, "1. Saludo y contexto", level=1)
    add_para(
        doc,
        "Estimado Antonio,",
    )
    add_para(
        doc,
        "Junto con saludarte, agradecemos tu interés en Imporlan. A continuación encontrarás "
        "la respuesta detallada a todas tus consultas respecto del velero J/29 del año 1987 "
        "publicado en YachtWorld (https://www.yachtworld.com/yacht/1987-j-boats-j-29-9841792/), "
        "incluyendo el costo total entregado en la V Región, el desglose de cada partida, los "
        "momentos en que se realiza cada pago, el alcance de nuestra inspección pre-compra y "
        "las garantías asociadas.",
    )

    add_heading(doc, "2. Resumen ejecutivo – J/29 1987", level=1)
    add_para(
        doc,
        "Considerando un valor de compra en USA de USD 10.000 (publicado), el costo total "
        "del proceso completo —compra, inspección, importación e internación entregado en la "
        "V Región— asciende a:",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CLP $32.697.440 (IVA incluido)")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL
    r.font.name = "Calibri"

    add_para(
        doc,
        "Este valor corresponde a un servicio All-Inclusive: incluye inspección pre-compra, "
        "negociación con el vendedor, compra, traslado terrestre en USA, flete marítimo, agente "
        "de aduanas, internación, gastos portuarios y entrega en V Región.",
    )

    add_heading(doc, "3. Cotización itemizada", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Ítem"
    hdr[1].text = "Monto (CLP)"
    hdr[2].text = "% del total"
    for c in hdr:
        shade(c, "0B3D91")
        set_cell_borders(c)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)
                run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows = [
        ("Valor lancha compra en USA (USD 10.000)", "$9.200.000", "28%"),
        ("Servicio All-Inclusive (incluye FEE Imporlan)", "$20.263.640", "62%"),
        ("IVA (sobre valor CIF)", "$3.233.800", "10%"),
        ("Impuesto al Lujo", "No aplica", "0%"),
    ]
    for label, monto, pct in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = monto
        row[2].text = pct
        for i, c in enumerate(row):
            set_cell_borders(c)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"
                if i > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_row = table.add_row().cells
    total_row[0].text = "TOTAL"
    total_row[1].text = "$32.697.440"
    total_row[2].text = "100%"
    for i, c in enumerate(total_row):
        shade(c, "E8EEF7")
        set_cell_borders(c)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = AZUL
                run.font.name = "Calibri"
            if i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_para(
        doc,
        "Tipo de cambio referencial USD 1 = CLP 920. La cotización se reajusta el día del pago "
        "al vendedor según tipo de cambio observado del Banco Central.",
        size=10,
        color=GRIS,
    )

    add_heading(doc, "4. Etapas del proceso y momentos de pago", level=1)
    add_para(
        doc,
        "Trabajamos con cuatro hitos de pago, alineados con el avance real de la operación. "
        "Esto te entrega control sobre cada etapa y nunca te pedimos el total por adelantado:",
    )

    pago_table = doc.add_table(rows=1, cols=4)
    pago_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = pago_table.rows[0].cells
    hdr[0].text = "Hito"
    hdr[1].text = "Cuándo"
    hdr[2].text = "Qué se paga"
    hdr[3].text = "Monto referencial"
    for c in hdr:
        shade(c, "0B3D91")
        set_cell_borders(c)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pagos = [
        (
            "1. Inspección pre-compra",
            "Antes de viajar a inspeccionar",
            "Honorarios del surveyor + traslados en USA",
            "USD 800 – USD 1.500\n(según ubicación)",
        ),
        (
            "2. Compra al vendedor",
            "Una vez aprobada la inspección y cerrada la negociación",
            "Valor de compra del velero en USA",
            "$9.200.000 (USD 10.000)",
        ),
        (
            "3. Logística e importación",
            "Al iniciar traslado en USA y embarque marítimo",
            "Flete USA, ocean freight, seguro, FEE Imporlan parcial",
            "≈ $12.000.000",
        ),
        (
            "4. Internación y entrega",
            "Al arribo a Chile, previo a retiro del puerto",
            "Aduana, IVA, gastos portuarios, traslado a V Región, FEE Imporlan saldo",
            "≈ $11.500.000",
        ),
    ]
    for hito, cuando, que, monto in pagos:
        row = pago_table.add_row().cells
        row[0].text = hito
        row[1].text = cuando
        row[2].text = que
        row[3].text = monto
        for c in row:
            set_cell_borders(c)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    doc.add_paragraph()
    add_para(
        doc,
        "Importante: el valor de la inspección (Hito 1) es independiente del proceso All-Inclusive "
        "y se paga directamente al surveyor. Si decides avanzar con la importación luego de una "
        "inspección satisfactoria, los hitos 2, 3 y 4 ya están considerados dentro del total "
        "cotizado (CLP $32.697.440).",
        size=10,
        color=GRIS,
    )

    add_heading(doc, "5. Inspección pre-compra: alcance, valor y garantías", level=1)

    add_heading(doc, "5.1 Valor de la inspección", level=2)
    add_para(
        doc,
        "Para un J/29 de 30 pies en costa este de USA, la inspección profesional realizada por "
        "un surveyor SAMS® / NAMS® acreditado tiene un valor referencial entre USD 800 y "
        "USD 1.500 (CLP $735.000 – $1.380.000 aprox.), dependiendo de la ubicación de la "
        "embarcación y de si se realiza prueba de mar (sea trial). En el caso del J/29 "
        "publicado, estimamos un valor cercano a USD 1.000 con prueba de mar incluida.",
    )

    add_heading(doc, "5.2 Qué incluye la inspección (detalle por sistema)", level=2)
    add_para(doc, "La inspección que coordinamos cubre los siguientes ítems:", bold=True)

    insp_items = [
        ("Electrónica e instrumentos: ",
         "GPS/plotter, VHF, ecosonda, anemómetro, corredera, piloto automático, batería de "
         "servicio y arranque, panel eléctrico, cargador, conexiones, cableado y puesta a tierra."),
        ("Motor: ",
         "estado general, horas de uso reales, compresión, fugas de aceite/refrigerante, sistema "
         "de enfriamiento, transmisión, hélice, eje, prensa-estopa, escape, filtros, alternador, "
         "y prueba de funcionamiento en marcha."),
        ("Casco, pintura y osmosis: ",
         "inspección visual exterior e interior, percusión del casco para detectar delaminaciones, "
         "test de humedad con higrómetro en obra viva, revisión de osmosis, antifouling, "
         "transducers, pasacascos, válvulas de fondo y zincos."),
        ("Velas y cabos: ",
         "inventario completo (mayor, génova, foques, spinnaker, asimétrico si aplica), revisión "
         "de costuras, refuerzos, baluma, grátil; cabos de maniobra (drizas, escotas, "
         "contras, amantillos), winches, mordazas y poleas."),
        ("Arboladura: ",
         "palo, botavara, tangón, crucetas, jarcia firme (stays popel y proel, obenques, "
         "burdas), terminales swage, uniones inox-aluminio (corrosión galvánica), pasadores, "
         "tornapuntas y herrajes."),
        ("Carro de arrastre o de patio: ",
         "estado del chasis, ejes, rodamientos, neumáticos, frenos, sistema eléctrico de luces, "
         "documentación y compatibilidad con normativa chilena para traslado en ruta."),
        ("Documentación y trazabilidad: ",
         "título (Title), HIN, mantenciones, historial de propietarios, deudas, embargos (UCC), "
         "registro USCG si aplica."),
    ]
    for prefix, body in insp_items:
        add_bullet(doc, body, bold_prefix=prefix)

    add_heading(doc, "5.3 Nivel de detalle", level=2)
    add_para(
        doc,
        "La inspección entrega un informe escrito (típicamente 30 a 60 páginas) con fotografías, "
        "lecturas instrumentales (humedad, compresión, voltajes), inventario completo, listado "
        "de hallazgos clasificados por criticidad (seguridad / funcional / cosmético) y "
        "recomendación de valor justo de mercado (fair market value). Si se contrata prueba de "
        "mar, también se evalúa rendimiento, vibraciones, manejo y comportamiento bajo carga.",
    )

    add_heading(doc, "5.4 Garantías sobre la inspección", level=2)
    add_para(
        doc,
        "Es importante ser transparentes con este punto: una inspección pre-compra es una "
        "fotografía técnica del estado del velero en una fecha determinada. Las garantías "
        "que entrega son:",
    )
    add_bullet(
        doc,
        "Responsabilidad profesional del surveyor: el inspector es un profesional acreditado "
        "(SAMS®/NAMS®) con seguro de responsabilidad civil (E&O). Si omite un defecto "
        "evidente o detectable mediante los métodos estándar de inspección, responde "
        "profesionalmente.",
        bold_prefix="• ",
    )
    add_bullet(
        doc,
        "Argumentos para renegociar precio: cualquier hallazgo permite ajustar el valor "
        "final con el vendedor o solicitar reparaciones antes del cierre.",
        bold_prefix="• ",
    )
    add_bullet(
        doc,
        "Decisión informada: si el informe arroja problemas mayores, puedes desistir de la "
        "compra sin haber comprometido el grueso del presupuesto (solo perderías el costo "
        "de la inspección, ~USD 1.000).",
        bold_prefix="• ",
    )
    add_para(
        doc,
        "Lo que la inspección NO garantiza: defectos ocultos no detectables sin desarme, "
        "fallas futuras del motor o equipos, ni el comportamiento del velero después de la "
        "compra. Por esto recomendamos siempre que, además de la inspección, se contrate un "
        "seguro de embarcación que cubra el período de traslado y posterior uso en Chile.",
        size=10,
        color=GRIS,
    )

    add_heading(doc, "6. Negociación con el vendedor", level=1)
    add_para(
        doc,
        "Sí, en Imporlan nos encargamos íntegramente de la negociación con el vendedor: tú no "
        "tienes que tratar directamente con el dueño ni preocuparte de transferencias "
        "internacionales. El proceso funciona así:",
    )
    add_bullet(
        doc,
        "Tomamos contacto con el broker o vendedor del J/29, validamos disponibilidad y "
        "negociamos precio y condiciones en tu nombre.",
        bold_prefix="1. ",
    )
    add_bullet(
        doc,
        "Coordinamos la inspección y, en base al informe, ajustamos la oferta final.",
        bold_prefix="2. ",
    )
    add_bullet(
        doc,
        "Firmamos el Purchase & Sale Agreement con cláusula de aprobación sujeta a inspección "
        "(estándar en USA) que protege tu depósito.",
        bold_prefix="3. ",
    )
    add_bullet(
        doc,
        "Realizamos el pago al vendedor mediante wire transfer desde nuestra estructura, contra "
        "entrega del título libre de gravámenes. Tú nos transfieres a nosotros en CLP.",
        bold_prefix="4. ",
    )
    add_bullet(
        doc,
        "Tomamos posesión, gestionamos el traslado terrestre al puerto de embarque y damos "
        "inicio al proceso de importación.",
        bold_prefix="5. ",
    )

    add_heading(doc, "7. Transparencia: FEE Imporlan", level=1)
    add_para(
        doc,
        "Por política de transparencia, te informamos desde el primer contacto cuál es nuestro "
        "honorario:",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FEE Imporlan: CLP $3.000.000")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = AZUL
    r.font.name = "Calibri"
    add_para(
        doc,
        "Este monto está contabilizado dentro del Servicio All-Inclusive (CLP $20.263.640) y "
        "corresponde a lo que efectivamente queda para Imporlan por gestionar la operación "
        "completa. El resto del All-Inclusive ($17.263.640) se destina íntegramente a "
        "gastos reales de la operación: flete USA, ocean freight, seguro, agente de aduanas, "
        "gastos portuarios, traslado a V Región y honorarios del surveyor (si se contrata "
        "junto con el paquete).",
    )

    add_heading(doc, "8. Próximos pasos sugeridos", level=1)
    add_bullet(
        doc,
        "Confirmar interés en avanzar con el J/29 1987 (link YachtWorld 9841792).",
        bold_prefix="1. ",
    )
    add_bullet(
        doc,
        "Validar disponibilidad y precio actual con el broker (lo hacemos nosotros en 24-48 h).",
        bold_prefix="2. ",
    )
    add_bullet(
        doc,
        "Coordinar y agendar la inspección pre-compra (pago Hito 1).",
        bold_prefix="3. ",
    )
    add_bullet(
        doc,
        "Revisar informe juntos y decidir si avanzamos a compra (pago Hito 2).",
        bold_prefix="4. ",
    )
    add_bullet(
        doc,
        "Activar logística de importación (pago Hito 3).",
        bold_prefix="5. ",
    )
    add_bullet(
        doc,
        "Internación, retiro de puerto y entrega en V Región (pago Hito 4).",
        bold_prefix="6. ",
    )

    add_para(
        doc,
        "Plazo total estimado del proceso: 10 a 14 semanas desde el cierre de compra hasta "
        "la entrega en V Región.",
        bold=True,
    )

    doc.add_paragraph()
    add_heading(doc, "Contacto", level=1)
    add_para(doc, "Imporlan – Importación de Embarcaciones", bold=True)
    add_para(doc, "Web: www.imporlan.cl")
    add_para(doc, "WhatsApp: +56 9 4021 1459")
    add_para(doc, "Email: contacto@imporlan.cl")

    doc.add_paragraph()
    closing = doc.add_paragraph()
    r = closing.add_run(
        "Quedamos atentos a tus comentarios para coordinar el primer paso. Si lo prefieres, "
        "podemos agendar una videollamada de 30 minutos para revisar este documento en "
        "conjunto."
    )
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GRIS
    r.font.name = "Calibri"

    output = "/home/user/Imporlan/docs/clientes/Respuesta_Antonio_J29_1987.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    path = build()
    print(f"Documento generado: {path}")
